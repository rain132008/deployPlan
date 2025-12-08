from docx import Document
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.path.join(BASE_DIR, "templates")
STEPS_DIR = os.path.join(TEMPLATES_DIR, "steps")

os.makedirs(STEPS_DIR, exist_ok=True)

def create_master():
    doc = Document()
    doc.add_heading('Release Plan: {{ version }}', 0)
    doc.add_paragraph('Owner: {{ owner }}')
    doc.add_paragraph('Date: {{ deploy_date }}')
    doc.add_heading('Deployment Steps', level=1)
    
    # Placeholder for steps. In docxtpl, to insert a list of subdocs, 
    # we usually just put a tag that the context will replace, but context replacement for subdocs 
    # works by inserting them at a location.
    # The generator code uses `doc.new_subdoc` and then `rendered_steps`.
    # Wait, in the generator code I wrote: `context['steps'] = rendered_steps`.
    # And `doc.render(context)`. 
    # In the template, we need `{{ steps }}`?
    # Actually, if we pass a LIST of subdocs to a tag, `docxtpl` doesn't automatically iterate and insert them unless we use a Jinja loop.
    # If `steps` is the list of subdoc objects:
    # {% for s in steps %}
    # {{ s }}
    # {% endfor %}
    
    p = doc.add_paragraph()
    runner = p.add_run('{% for s in steps %}{{ s }}\n{% endfor %}')
    
    path = os.path.join(TEMPLATES_DIR, "master.docx")
    doc.save(path)
    print(f"Created {path}")

def create_java_step():
    doc = Document()
    doc.add_heading('Java Deployment', level=2)
    
    # Docxtpl table loop syntax:
    # We want a row for each item in the 'items' list.
    # The columns will be: App Name, Unit, Version
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'App Name'
    hdr_cells[1].text = 'Unit'
    hdr_cells[2].text = 'Version'
    
    # Data Row (Template)
    # We add a row with variables, and put the loop tags around the row or inside the first/last cell depending on complexity.
    # For docxtpl, standard way is:
    # {% tr for item in items %} ... cell contents ... {% tr endfor %}
    # But python-docx doesn't let us easily write raw XML tags around rows.
    # Docxtpl lets us put `{% tr for item in items %}` in the first cell of the row, 
    # and `{% tr endfor %}` in the last cell (or implicit if same row).
    # Ideally:
    # Cell 1: {% tr for item in items %}{{ item.app_name }}
    # Cell 2: {{ item.deploy_unit }}
    # Cell 3: {{ item.pkg_version }}{% tr endfor %}
    
    row_cells = table.add_row().cells
    row_cells[0].text = '{% tr for item in items %}{{ item.app_name }}'
    row_cells[1].text = '{{ item.deploy_unit }}'
    row_cells[2].text = '{{ item.pkg_version }}{% tr endfor %}'
    
    path = os.path.join(STEPS_DIR, "step_java.docx")
    doc.save(path)
    print(f"Created {path}")

def create_sql_step():
    doc = Document()
    doc.add_heading('SQL Execution', level=2)
    p = doc.add_paragraph()
    p.add_run('DB Code: {{ db_name }}\n')
    p.add_run('Script Content:\n{{ sql_content }}')
    
    path = os.path.join(STEPS_DIR, "step_sql.docx")
    doc.save(path)
    print(f"Created {path}")

if __name__ == "__main__":
    create_master()
    create_java_step()
    create_sql_step()
